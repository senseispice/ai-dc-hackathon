"""Tests for app/tools/document_parser.py"""

from __future__ import annotations

import pathlib
import tempfile
import pytest

from app.tools.document_parser import parse_document, parse_docx, parse_pdf, chunk_text

DATA_DIR = pathlib.Path(__file__).parent.parent / "data"


# ---------------------------------------------------------------------------
# parse_document — .txt
# ---------------------------------------------------------------------------

def test_parse_document_txt_returns_string():
    text = parse_document(str(DATA_DIR / "sample_rfp.txt"))
    assert isinstance(text, str)


def test_parse_document_txt_not_empty():
    text = parse_document(str(DATA_DIR / "sample_rfp.txt"))
    assert len(text) > 0


def test_parse_document_txt_content():
    """Check that known content from the file actually appears in output."""
    text = parse_document(str(DATA_DIR / "sample_rfp.txt"))
    # sample_rfp.txt is a BLM solicitation document
    assert "BLM" in text or "REQUEST FOR PROPOSALS" in text


def test_parse_document_unsupported_extension_raises():
    with pytest.raises(ValueError, match="Unsupported file type"):
        parse_document("somefile.xyz")


# ---------------------------------------------------------------------------
# parse_document — .pdf (written to a temp file)
# ---------------------------------------------------------------------------

def test_parse_pdf_returns_string():
    fitz = pytest.importorskip("fitz")
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
        # Create a minimal single-page PDF with known text
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "Hello PDF world")
        doc.save(f.name)
        doc.close()
        tmp_path = f.name

    text = parse_pdf(tmp_path)
    assert isinstance(text, str)
    assert "Hello PDF world" in text


def test_parse_document_routes_pdf():
    fitz = pytest.importorskip("fitz")
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "PDF routing test")
        doc.save(f.name)
        doc.close()
        tmp_path = f.name

    text = parse_document(tmp_path)
    assert "PDF routing test" in text


# ---------------------------------------------------------------------------
# parse_document — .docx (written to a temp file)
# ---------------------------------------------------------------------------

def test_parse_docx_returns_string():
    docx = pytest.importorskip("docx")
    from docx import Document

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        tmp_path = f.name

    doc = Document()
    doc.add_paragraph("Hello DOCX world")
    doc.save(tmp_path)

    text = parse_docx(tmp_path)
    assert isinstance(text, str)
    assert "Hello DOCX world" in text


def test_parse_document_routes_docx():
    pytest.importorskip("docx")
    from docx import Document

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        tmp_path = f.name

    doc = Document()
    doc.add_paragraph("DOCX routing test")
    doc.save(tmp_path)

    text = parse_document(tmp_path)
    assert "DOCX routing test" in text


# ---------------------------------------------------------------------------
# chunk_text
# ---------------------------------------------------------------------------

def test_chunk_text_returns_list():
    chunks = chunk_text("hello world")
    assert isinstance(chunks, list)


def test_chunk_text_single_chunk_when_small():
    text = "Short text that fits in one chunk."
    chunks = chunk_text(text, chunk_size=2000, overlap=200)
    assert len(chunks) == 1


def test_chunk_text_multiple_chunks():
    # Generate text that is definitely larger than 10 tokens
    text = " ".join(["word"] * 50)
    chunks = chunk_text(text, chunk_size=10, overlap=2)
    assert len(chunks) > 1


def test_chunk_text_overlap_means_shared_content():
    """With overlap, consecutive chunks should share tokens at the boundary."""
    text = " ".join([str(i) for i in range(200)])
    chunks = chunk_text(text, chunk_size=20, overlap=5)
    # The tail of chunk N should appear at the head of chunk N+1
    assert len(chunks) >= 2
    tail = chunks[0].split()[-3:]
    head = chunks[1].split()[:3]
    # At least some of the tail words should reappear at the start of the next chunk
    assert any(w in head for w in tail)


def test_chunk_text_reconstructs_all_content():
    """Every token in the original should appear in at least one chunk."""
    text = "The quick brown fox jumps over the lazy dog"
    chunks = chunk_text(text, chunk_size=5, overlap=1)
    combined = " ".join(chunks)
    for word in text.split():
        assert word in combined


def test_chunk_text_no_empty_chunks():
    text = "Some reasonable body of text for testing purposes."
    chunks = chunk_text(text)
    assert all(len(c) > 0 for c in chunks)
